#Import necessary libraries
from docx import Document #docx manipulation library
from sys import argv
import os.path
from datetime import datetime

import boto3 #AWS Python SDK library
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
ses = boto3.client('ses')


script, req_date,req_time,req_client,req_department,req_address,req_phone,req_contact,req_system,type,req_desc,req_nameof,engineer = argv

template = Document('S_R_F.docx')

newFile = template
tab = newFile.tables[0]


tab.cell(0,1).paragraphs[0].clear()
tab.cell(0,1).paragraphs[0].text = req_date

tab.cell(1,1).paragraphs[0].clear()
tab.cell(1,1).paragraphs[0].text = req_time

tab.cell(2,1).paragraphs[0].clear()
tab.cell(2,1).paragraphs[0].text = req_client

tab.cell(3,1).paragraphs[0].clear()
tab.cell(3,1).paragraphs[0].text = req_department

tab.cell(4,1).paragraphs[0].clear()
tab.cell(4,1).paragraphs[0].text = req_address

tab.cell(5,1).paragraphs[0].clear()
tab.cell(5,1).paragraphs[0].text = req_phone

tab.cell(6,1).paragraphs[0].clear()
tab.cell(6,1).paragraphs[0].text = req_contact

tab.cell(7,1).paragraphs[0].clear()
tab.cell(7,1).paragraphs[0].text = req_system

tab.cell(8,1).paragraphs[0].clear()
tab.cell(8,1).paragraphs[0].text = type

tab.cell(9,1).paragraphs[0].clear()
tab.cell(9,1).paragraphs[0].text = req_desc

newFileName = req_client+'_'+datetime.now().strftime("%Y-%m-%d %H:%M:%S")
newFileName = 'S_Req_'+newFileName.replace(" ", "_")+'.docx'
newFile.save(newFileName)
#newFile.close()

"""
This section is for sending the email using AWS raw email via SES
"""

message = MIMEMultipart()
message['Subject'] = 'Incoming Service Request'
message['From'] = 'CCE-Service <service@cce-service.ca>'
message['To'] = ', '.join(['sbrukson@gmail.com', 'alex@ccemedical.com'])
# message body
body_string = 'Submitted By: '+req_nameof+'<br/> Requested Engineer: '+engineer+' <br/><br/>'
body_string += 'Date: '+req_date+'<br/> Client: '+req_client+' <br/> Address: '+req_address
body_string += '<br/> Systems: '+req_system+'<br/> Description: '+req_desc


part = MIMEText(body_string, 'html')
message.attach(part)
# attachment

part = MIMEApplication(open(newFileName, 'rb').read())

part.add_header('Content-Disposition', 'attachment', filename=newFileName)
message.attach(part)
response = ses.send_raw_email(
    Source=message['From'],
    Destinations=['sbrukson@gmail.com', 'alex@ccemedical.com'],
    RawMessage={
        'Data': message.as_string()
    }
)

print(response)
